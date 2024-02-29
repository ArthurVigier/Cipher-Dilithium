from flask import Flask, request, jsonify
import os
import base64
import io
import uuid
import ast
import json
from datetime import datetime
from API_DIRECTORY.dilithium import Dilithium
import zipfile
from PyPDF2 import PdfReader
from docx import Document
import pickle
import logging
app = Flask(__name__)

if __name__ != "__main__":
    gunicorn_logger = logging.getLogger('gunicorn.error')
    app.logger.handlers = gunicorn_logger.handlers
    app.logger.setLevel(gunicorn_logger.level)

# Durée de l'enregistrement en secondes
RECORD_SECONDS = 5
# Fréquence d'échantillonnage
SAMPLE_RATE = 44100
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.asymmetric import ec
from cryptography.hazmat.primitives import serialization
from cryptography.exceptions import InvalidSignature

def generate_signature(message):
    # Générer une clé privée
    private_key = ec.generate_private_key(ec.SECP256R1())

    # Signer le message
    ecc_signature = private_key.sign(message.encode('utf-8'), ec.ECDSA(hashes.SHA256()))

    # Obtenir la clé publique
    ecc_public_key = private_key.public_key().public_bytes(
        encoding=serialization.Encoding.PEM,
        format=serialization.PublicFormat.SubjectPublicKeyInfo
    )

    # Dilithium signature generation
    dilithium_signer = Dilithium()  # Create a Dilithium signer instance
    pq_signature = dilithium_signer.sign(message)

    # Conserver la clé publique Dilithium sous sa forme originale
    pq_public_key = dilithium_signer.pk

    print(f"pq_signature format: {type(pq_signature)}")
    print(f"pq_public_key format: {type(pq_public_key)}")
    return pq_signature, ecc_signature, pq_public_key, ecc_public_key



def verify_signature(message, pq_signature, pq_public_key , ecc_signature, ecc_public_key):
    # Vérifier si le message est fourni
    if not message:
        raise ValueError("Le message est requis pour la vérification de la signature.")
    # Afficher les valeurs des signatures et des clés publiques

    # Initialiser les vérificateurs
    dilithium_verifier = Dilithium()
    ecc_verifier = ec.ECDSA(hashes.SHA256())

    # Créer un dictionnaire pour stocker les résultats
    results = {'dilithium': {'success': None, 'error': None}, 'ecc': {'success': None, 'error': None}}

    # Tenter la vérification Dilithium
    try:
        print("Début de la vérification Dilithium")
        pk = pickle.loads(pq_public_key)
        print("Clé publique Dilithium chargée")
        results['dilithium']['success'] = dilithium_verifier.verify(message, pickle.loads(pq_signature), pk)
        print("Vérification Dilithium terminée")
    except Exception as e:
        print(f"Erreur lors de la vérification Dilithium : {e}")
        results['dilithium']['success'] = False
        results['dilithium']['error'] = str(e)

    # Tenter la vérification ECC
    try:
        print("Début de la vérification ECC")
        public_key = serialization.load_pem_public_key(ecc_public_key)
        print("Clé publique ECC chargée")
        public_key.verify(ecc_signature, message, ecc_verifier)
        print("Vérification ECC terminée")
        results['ecc']['success'] = True
    except InvalidSignature as e:
        print("Erreur lors de la vérification ECC : signature invalide")
        results['ecc']['success'] = False
        results['ecc']['error'] = 'Signature invalide'

    return results


@app.route('/generate_signature', methods=['POST'])
def generate_signature_route():
    message = request.json.get('message')
    if not message:
        return jsonify({"error": "No message provided"}), 400

    # Convert the message to bytes if it's not already
    if isinstance(message, str):
        message = message.encode()

    pq_signature_str, ecc_signature, pq_public_key_str, ecc_public_key = generate_signature(message)

    return jsonify({
        'pq_signature_str': pq_signature_str,
        'ecc_signature': ecc_signature,
        'public_key_str': pq_public_key_str,
        'ecc_public_key': ecc_public_key,
        'message': 'Signature generated successfully',
    })




@app.route('/generate_signature_from_file', methods=['POST'])
def generate_signature_from_file_route():
    # Check if a file was provided
    if 'file' not in request.files:
        return "No file provided", 400

    file = request.files['file']
    file_extension = file.filename.rsplit('.', 1)[1].lower()

    # Read the file content based on its type
    if file_extension == 'pdf':
        reader = PdfReader(file.stream)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    elif file_extension == 'docx':
        doc = Document(file)
        text = " ".join([p.text for p in doc.paragraphs])
    else:
        text = file.read().decode('utf-8')

    # Use the text as the message for generate_signature
    pq_signature, ecc_signature, pq_public_key, ecc_public_key = generate_signature(text)

    # Prepare the data to be written to the file
    data = {
        'message': text,
        'pq_signature': pq_signature,
        'ecc_signature': ecc_signature,
        'pq_public_key': pq_public_key,
        'ecc_public_key': ecc_public_key,
        'date': datetime.now().isoformat()  # Add the current date and time
    }

    file_name = 'signature.zip'

    # Write the data to a zip file
    with zipfile.ZipFile(file_name, 'w') as zipf:
        for key, value in data.items():
            if key == 'ecc_public_key':
                # Keep ecc_public_key in PEM format and write it to a .txt file
                zipf.writestr(f'{key}.txt', value.decode('utf-8'))
            elif key in ['pq_signature', 'pq_public_key', 'message', 'ecc_signature']:  # Add 'ecc_signature' here
                # Convert tuples to bytes and keep 'message' and 'ecc_signature' as bytes
                if key not in ['message', 'ecc_signature']:  # Exclude 'message' and 'ecc_signature' from conversion
                    value = pickle.dumps(value)
                zipf.writestr(f'{key}.bin', value)
            else:
                # Convert other values to strings
                value = str(value).encode('utf-8')
                zipf.writestr(f'{key}.bin', value)

    # Print the file path
    print(f'File saved at: {file_name}')

    return jsonify(data)

@app.route('/verify_signature', methods=['POST'])
def verify_signature_route():
    # Check if a file was provided
    if 'file' not in request.files or 'file_to_verify' not in request.files:
        return jsonify({"error": "No file or file_to_verify provided"}), 400

    file = request.files['file']
    file_to_verify = request.files['file_to_verify']
    file_extension = file_to_verify.filename.rsplit('.', 1)[1].lower()

    # Read the file content based on its type
    if file_extension == 'pdf':
        reader = PdfReader(file_to_verify.stream)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    elif file_extension == 'docx':
        doc = Document(file_to_verify)
        text = " ".join([p.text for p in doc.paragraphs])
    else:
        text = file_to_verify.read().decode('utf-8')

    app.logger.info('File and text received for signature verification')

    # Read the file data
    zip_file = zipfile.ZipFile(io.BytesIO(file.read()))

    data = {}
    for filename in zip_file.namelist():
        with zip_file.open(filename) as f:
            key = filename.rsplit('.', 1)[0]
            value = f.read()
            data[key] = value

    # Compare the text with the one from the zip file
    if text != data['message']:
        return jsonify({'result comparison': False})

    app.logger.info('Comparison successful')

    # Call the verify_signature function with the original data
    result = verify_signature(
        data['message'], data['pq_signature'], data['pq_public_key'], data['ecc_signature'], data['ecc_public_key']
    )

    # Return the result as a JSON response
    return jsonify({'result': result})
if __name__ == '__main__':
    app.run(debug=True)